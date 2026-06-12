"""File parsers for supported document types."""
import io
from typing import Tuple

from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from markdown_it import MarkdownIt


class UnsupportedFileType(Exception):
    pass


def _ext(name: str) -> str:
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def parse_file(filename: str, content: bytes) -> Tuple[str, str]:
    """Return (extracted_text, mime_hint)."""
    ext = _ext(filename)

    if ext == "pdf":
        reader = PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(parts), "application/pdf"

    if ext == "docx":
        doc = DocxDocument(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if ext == "doc":
        raise UnsupportedFileType(
            "Legacy .doc files aren't supported yet — please convert to .docx first."
        )

    if ext in ("html", "htm"):
        soup = BeautifulSoup(content.decode("utf-8", errors="ignore"), "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text(separator="\n"), "text/html"

    if ext in ("md", "markdown"):
        md = MarkdownIt()
        text = content.decode("utf-8", errors="ignore")
        tokens = md.parse(text)
        # Crude: just return raw markdown text — it's already mostly plain
        return text, "text/markdown"

    if ext == "txt" or ext == "":
        return content.decode("utf-8", errors="ignore"), "text/plain"

    raise UnsupportedFileType(f"Unsupported file extension: .{ext}")
